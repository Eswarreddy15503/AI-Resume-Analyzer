import io
import docx
import logging

logger = logging.getLogger(__name__)

def extract_text_from_docx(docx_file) -> str:
    """
    Extracts text from a DOCX file stream or file path.
    
    Args:
        docx_file: A file-like object (e.g., BytesIO from Streamlit) or a file path.
        
    Returns:
        str: Extracted text from the DOCX.
    """
    text = []
    try:
        # If docx_file is bytes, wrap it in BytesIO
        if isinstance(docx_file, bytes):
            docx_file = io.BytesIO(docx_file)
            
        doc = docx.Document(docx_file)
        
        # 1. Extract from paragraphs
        for para in doc.paragraphs:
            if para.text:
                text.append(para.text)
                
        # 2. Extract from tables (resumes often use tables for layout)
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text and cell_text not in row_text:
                        row_text.append(cell_text)
                if row_text:
                    text.append(" | ".join(row_text))
                    
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}")
        raise RuntimeError(f"Could not parse DOCX file: {str(e)}")
        
    return "\n".join(text).strip()

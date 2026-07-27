import os
import sys

# Define sample content
RESUME_TEXT = """JOHN DOE
Email: john.doe@email.com | Phone: (123) 456-7890
LinkedIn: linkedin.com/in/johndoe | GitHub: github.com/johndoe

PROFESSIONAL SUMMARY
Results-driven Software Engineer with over 4 years of experience designing and deploying scalable web applications. Proven expertise in cloud migrations and API optimization.

WORK EXPERIENCE
Senior Developer | Tech Solutions Inc. | 2022 - Present
- Led the redevelopment of a legacy backend services system, migration to microservices, and deployment to AWS.
- Designed and built REST APIs using Python, FastAPI, and PostgreSQL, which optimized application response times by 35%.
- Automated CI/CD deployment pipelines using GitHub Actions and Docker, reducing deployment times by 40%.
- Managed and mentored a team of 3 junior developers on Agile practices and Git workflows.

Software Engineer | DevCorp | 2020 - 2022
- Engineered web services using Java, Spring Boot, and MySQL databases.
- Collaborated with product managers to deliver 5 high-priority feature integrations.
- Integrated Prometheus and Grafana monitoring tools to track infrastructure health.

PROJECTS
E-Commerce API Service (Personal Project)
- Developed a high-performance shopping cart backend using Go, Redis, and MongoDB.
- Containerized application services using Docker and deployed to AWS ECS.

EDUCATION
Bachelor of Science in Computer Science
University of Technology | 2016 - 2020

CERTIFICATIONS
- AWS Certified Solutions Architect - Associate
- Certified ScrumMaster (CSM)

SKILLS
- Programming Languages: Python, Java, Go, SQL, HTML, CSS, JavaScript
- Frameworks: FastAPI, Django, Spring Boot, React
- Databases: PostgreSQL, MySQL, Redis, MongoDB
- Cloud & DevOps: AWS, Docker, GitHub Actions, CI/CD, Prometheus, Grafana
- Tools: Git, GitHub, Jira, VS Code
- Soft Skills: Leadership, Collaboration, Agile, Problem Solving, Communication
"""

def make_docx():
    try:
        import docx
        doc = docx.Document()
        doc.add_heading('John Doe', 0)
        
        doc.add_paragraph('Email: john.doe@email.com | Phone: (123) 456-7890\nLinkedIn: linkedin.com/in/johndoe | GitHub: github.com/johndoe')
        
        doc.add_heading('Professional Summary', level=1)
        doc.add_paragraph('Results-driven Software Engineer with over 4 years of experience designing and deploying scalable web applications. Proven expertise in cloud migrations and API optimization.')
        
        doc.add_heading('Work Experience', level=1)
        doc.add_heading('Senior Developer | Tech Solutions Inc. | 2022 - Present', level=2)
        doc.add_paragraph('- Led the redevelopment of a legacy backend services system, migration to microservices, and deployment to AWS.')
        doc.add_paragraph('- Designed and built REST APIs using Python, FastAPI, and PostgreSQL, which optimized application response times by 35%.')
        doc.add_paragraph('- Automated CI/CD deployment pipelines using GitHub Actions and Docker, reducing deployment times by 40%.')
        doc.add_paragraph('- Managed and mentored a team of 3 junior developers on Agile practices and Git workflows.')
        
        doc.add_heading('Software Engineer | DevCorp | 2020 - 2022', level=2)
        doc.add_paragraph('- Engineered web services using Java, Spring Boot, and MySQL databases.')
        doc.add_paragraph('- Collaborated with product managers to deliver 5 high-priority feature integrations.')
        doc.add_paragraph('- Integrated Prometheus and Grafana monitoring tools to track infrastructure health.')
        
        doc.add_heading('Projects', level=1)
        doc.add_heading('E-Commerce API Service (Personal Project)', level=2)
        doc.add_paragraph('- Developed a high-performance shopping cart backend using Go, Redis, and MongoDB.')
        doc.add_paragraph('- Containerized application services using Docker and deployed to AWS ECS.')
        
        doc.add_heading('Education', level=1)
        doc.add_paragraph('Bachelor of Science in Computer Science\nUniversity of Technology | 2016 - 2020')
        
        doc.add_heading('Certifications', level=1)
        doc.add_paragraph('- AWS Certified Solutions Architect - Associate\n- Certified ScrumMaster (CSM)')
        
        doc.add_heading('Skills', level=1)
        doc.add_paragraph('Programming Languages: Python, Java, Go, SQL, HTML, CSS, JavaScript\n'
                          'Frameworks: FastAPI, Django, Spring Boot, React\n'
                          'Databases: PostgreSQL, MySQL, Redis, MongoDB\n'
                          'Cloud & DevOps: AWS, Docker, GitHub Actions, CI/CD, Prometheus, Grafana\n'
                          'Tools: Git, GitHub, Jira, VS Code\n'
                          'Soft Skills: Leadership, Collaboration, Agile, Problem Solving, Communication')
        
        doc.save('sample_resume.docx')
        print("Successfully generated sample_resume.docx")
    except Exception as e:
        print(f"Error creating DOCX: {e}")

def make_pdf():
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        
        doc = SimpleDocTemplate("sample_resume.pdf", pagesize=letter, rightMargin=50, leftMargin=50, topMargin=50, bottomMargin=50)
        styles = getSampleStyleSheet()
        
        story = []
        
        # Style variations
        title_style = ParagraphStyle('T', fontName='Helvetica-Bold', fontSize=22, spaceAfter=4, textColor=colors.HexColor("#1E3A8A"))
        sub_style = ParagraphStyle('S', fontName='Helvetica', fontSize=10, spaceAfter=15, textColor=colors.HexColor("#4B5563"))
        h1_style = ParagraphStyle('H1', fontName='Helvetica-Bold', fontSize=14, spaceBefore=10, spaceAfter=5, textColor=colors.HexColor("#1E3A8A"))
        h2_style = ParagraphStyle('H2', fontName='Helvetica-Bold', fontSize=11, spaceBefore=6, spaceAfter=2, textColor=colors.HexColor("#3B82F6"))
        body_style = ParagraphStyle('B', fontName='Helvetica', fontSize=10, leading=14, spaceAfter=4)
        bullet_style = ParagraphStyle('Bu', fontName='Helvetica', fontSize=10, leading=13, leftIndent=15, firstLineIndent=-10, spaceAfter=3)
        
        story.append(Paragraph("JOHN DOE", title_style))
        story.append(Paragraph("Email: john.doe@email.com | Phone: (123) 456-7890<br/>LinkedIn: linkedin.com/in/johndoe | GitHub: github.com/johndoe", sub_style))
        
        story.append(Paragraph("PROFESSIONAL SUMMARY", h1_style))
        story.append(Paragraph("Results-driven Software Engineer with over 4 years of experience designing and deploying scalable web applications. Proven expertise in cloud migrations and API optimization.", body_style))
        
        story.append(Paragraph("WORK EXPERIENCE", h1_style))
        story.append(Paragraph("Senior Developer | Tech Solutions Inc. | 2022 - Present", h2_style))
        story.append(Paragraph("• Led the redevelopment of a legacy backend services system, migration to microservices, and deployment to AWS.", bullet_style))
        story.append(Paragraph("• Designed and built REST APIs using Python, FastAPI, and PostgreSQL, which optimized application response times by 35%.", bullet_style))
        story.append(Paragraph("• Automated CI/CD deployment pipelines using GitHub Actions and Docker, reducing deployment times by 40%.", bullet_style))
        story.append(Paragraph("• Managed and mentored a team of 3 junior developers on Agile practices and Git workflows.", bullet_style))
        
        story.append(Paragraph("Software Engineer | DevCorp | 2020 - 2022", h2_style))
        story.append(Paragraph("• Engineered web services using Java, Spring Boot, and MySQL databases.", bullet_style))
        story.append(Paragraph("• Collaborated with product managers to deliver 5 high-priority feature integrations.", bullet_style))
        story.append(Paragraph("• Integrated Prometheus and Grafana monitoring tools to track infrastructure health.", bullet_style))
        
        story.append(Paragraph("PROJECTS", h1_style))
        story.append(Paragraph("E-Commerce API Service (Personal Project)", h2_style))
        story.append(Paragraph("• Developed a high-performance shopping cart backend using Go, Redis, and MongoDB.", bullet_style))
        story.append(Paragraph("• Containerized application services using Docker and deployed to AWS ECS.", bullet_style))
        
        story.append(Paragraph("EDUCATION", h1_style))
        story.append(Paragraph("Bachelor of Science in Computer Science<br/>University of Technology | 2016 - 2020", body_style))
        
        story.append(Paragraph("CERTIFICATIONS", h1_style))
        story.append(Paragraph("• AWS Certified Solutions Architect - Associate", bullet_style))
        story.append(Paragraph("• Certified ScrumMaster (CSM)", bullet_style))
        
        story.append(Paragraph("SKILLS", h1_style))
        story.append(Paragraph("<b>Programming Languages:</b> Python, Java, Go, SQL, HTML, CSS, JavaScript<br/>"
                               "<b>Frameworks:</b> FastAPI, Django, Spring Boot, React<br/>"
                               "<b>Databases:</b> PostgreSQL, MySQL, Redis, MongoDB<br/>"
                               "<b>Cloud & DevOps:</b> AWS, Docker, GitHub Actions, CI/CD, Prometheus, Grafana<br/>"
                               "<b>Tools:</b> Git, GitHub, Jira, VS Code<br/>"
                               "<b>Soft Skills:</b> Leadership, Collaboration, Agile, Problem Solving, Communication", body_style))
        
        doc.build(story)
        print("Successfully generated sample_resume.pdf")
    except Exception as e:
        print(f"Error creating PDF: {e}")

if __name__ == "__main__":
    make_docx()
    make_pdf()
